# -*- coding: utf-8 -*-
from chinaesi import main as i_
from chinanecc import main as c_
from chinappia import main as a_
from cncecc import main as cc_
from huanbao import main as o_
from syhgjn import main as n_
from docx import Document
import re

if __name__ == "__main__":
    sum_list = list()
    document = Document()
    for i in [i_, c_, a_, cc_, o_, n_]:
        for j in i():
            paragraph = document.add_paragraph()
            arial_run = paragraph.add_run(re.sub(r'\n+','\n',j.encode('utf-8').decode('utf-8'))
                                          + '\n'+'|'*78
                                          + '\n'+'|'*78)
            arial_run.font.name = 'SimSun'
            document.add_page_break()
    document.save('sum.docx')
